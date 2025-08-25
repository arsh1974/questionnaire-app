import os
import time
import random
import streamlit as st
import PyPDF2
from openai import OpenAI, error as openai_error
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# -------------------
# Setup OpenAI client
# -------------------
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("❌ No OpenAI API key found. Please set it in Streamlit Cloud Secrets.")
    st.stop()

client = OpenAI(api_key=api_key)

# -------------------
# Helper: Extract text from PDF
# -------------------
def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

# -------------------
# Helper: Generate Questions with Retry Logic
# -------------------
def generate_questions(text, level, model_choice):
    prompt = f"""
    You are a teacher creating a questionnaire for students.

    Content:
    {text[:3000]}

    Create 5 {level}-level questions based on this chapter.
    Format:
    Q1.
    Q2.
    Q3.
    Q4.
    Q5.
    """

    for attempt in range(5):  # retry up to 5 times
        try:
            response = client.chat.completions.create(
                model=model_choice,
                messages=[
                    {"role": "system", "content": "You are a helpful teacher."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            return response.choices[0].message.content.strip()

        except openai_error.RateLimitError:
            wait_time = (2 ** attempt) + random.random()
            st.warning(f"⏳ Rate limit hit. Retrying in {wait_time:.1f}s...")
            time.sleep(wait_time)

        except openai_error.OpenAIError as e:
            st.error(f"❌ OpenAI error: {str(e)}")
            return None

    return "❌ Failed after multiple retries. Please try again later."

# -------------------
# Helper: Save Word
# -------------------
def save_as_word(questions):
    doc = Document()
    doc.add_heading("Generated Questionnaire", 0)
    for q in questions.split("\n"):
        doc.add_paragraph(q)
    doc_path = "questionnaire.docx"
    doc.save(doc_path)
    return doc_path

# -------------------
# Helper: Save PDF
# -------------------
def save_as_pdf(questions):
    pdf_path = "questionnaire.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Times-Roman", 12)
    for line in questions.split("\n"):
        c.drawString(50, y, line)
        y -= 20
        if y < 50:
            c.showPage()
            c.setFont("Times-Roman", 12)
            y = height - 50
    c.save()
    return pdf_path

# -------------------
# Streamlit UI
# -------------------
st.title("📘 Questionnaire Generator for Teachers")

uploaded_file = st.file_uploader("Upload a PDF chapter", type="pdf")
level = st.selectbox("Choose difficulty level", ["Basic", "Advanced", "Expert"])
model_choice = st.radio(
    "Choose OpenAI model",
    ["gpt-4o-mini (recommended)", "gpt-3.5-turbo"]
)

if uploaded_file is not None:
    text = extract_text_from_pdf(uploaded_file)

    if st.button("Generate Questionnaire"):
        with st.spinner("Generating questions..."):
            questions = generate_questions(text, level, model_choice.split()[0])

        if questions:
            st.subheader("Generated Questionnaire")
            st.text_area("Questions", questions, height=250)

            # Download options
            word_file = save_as_word(questions)
            pdf_file = save_as_pdf(questions)

            with open(word_file, "rb") as f:
                st.download_button("⬇️ Download as Word", f, file_name="questionnaire.docx")

            with open(pdf_file, "rb") as f:
                st.download_button("⬇️ Download as PDF", f, file_name="questionnaire.pdf")
